# docx_to_txt.py
# pip install python-docx
# 需要安装antiword: sudo apt-get install antiword
import os, re, argparse, subprocess
from docx import Document

def normalize(s: str) -> str:
    s = s.replace("\u3000", " ").replace("\xa0", " ")  # 全角空格/nbsp
    s = re.sub(r"[ \t]+", " ", s)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # 把 "第  一  条 / 第 1 条 / 第一百二十条 之 一" 收敛为 "第…条"
    s = re.sub(r"第\s*([0-9零一二三四五六七八九十百千两]+)\s*条", r"第\1条", s)
    s = re.sub(r"(第[0-9零一二三四五六七八九十百千两]+条)\s*之\s*([一二三四五六七八九十]+)", r"\1之\2", s)
    return s.strip()

def docx_to_text(path: str) -> str:
    try:
        doc = Document(path)
    except Exception as e:
        # 关键：打印出错的 .docx 文件（你就能直接删/替换它）
        print(f"[ERROR] 解析 .docx 失败: {path}\n  -> {type(e).__name__}: {e}")
        return ""  # 返回空串，外层会跳过写出
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        lines.append(t if t else "")
    for tb in doc.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    text = "\n".join(lines)
    return normalize(text)

def doc_to_text(path: str) -> str:
    """使用antiword直接提取.doc文件内容"""
    try:
        # 使用antiword提取文本
        result = subprocess.run(['antiword', '-m', 'UTF-8.txt', path], 
                              stdout=subprocess.PIPE, stderr=subprocess.PIPE, 
                              text=True, encoding='utf-8')
        if result.returncode == 0:
            return normalize(result.stdout)
        else:
            print(f"[警告] antiword处理失败: {path}, 错误: {result.stderr}")
            return ""
    except FileNotFoundError:
        print("[错误] antiword未安装，请运行: sudo apt-get install antiword")
        return ""
    except Exception as e:
        print(f"[警告] 无法处理文件: {path}, 错误: {str(e)}")
        return ""

def process_directory(in_dir, out_dir, recursive=True):
    """处理目录中的所有.doc和.docx文件"""
    os.makedirs(out_dir, exist_ok=True)
    bad_files = []  # 新增：记录失败文件

    if recursive:
        for root, dirs, files in os.walk(in_dir):
            rel_path = os.path.relpath(root, in_dir)
            current_out_dir = out_dir if rel_path == '.' else os.path.join(out_dir, rel_path)
            os.makedirs(current_out_dir, exist_ok=True)

            for fn in files:
                if fn.lower().endswith(('.doc', '.docx')):
                    inp = os.path.join(root, fn)
                    out = os.path.join(current_out_dir, os.path.splitext(fn)[0] + ".txt")
                    try:
                        if fn.lower().endswith('.docx'):
                            txt = docx_to_text(inp)
                        else:
                            txt = doc_to_text(inp)
                    except Exception as e:
                        # 兜底打印，防止任何未预料异常中断批处理
                        print(f"[ERROR] 提取失败: {inp}\n  -> {type(e).__name__}: {e}")
                        bad_files.append(inp)
                        continue

                    if txt:
                        with open(out, "w", encoding="utf-8") as w:
                            w.write(txt)
                        print(f"[OK] {fn} -> {out} ({len(txt)} chars)")
                    else:
                        # docx_to_text/doc_to_text 返回空，也归入失败清单
                        bad_files.append(inp)
    else:
        for fn in os.listdir(in_dir):
            if fn.lower().endswith(('.doc', '.docx')):
                inp = os.path.join(in_dir, fn)
                out = os.path.join(out_dir, os.path.splitext(fn)[0] + ".txt")
                try:
                    if fn.lower().endswith('.docx'):
                        txt = docx_to_text(inp)
                    else:
                        txt = doc_to_text(inp)
                except Exception as e:
                    print(f"[ERROR] 提取失败: {inp}\n  -> {type(e).__name__}: {e}")
                    bad_files.append(inp)
                    continue

                if txt:
                    with open(out, "w", encoding="utf-8") as w:
                        w.write(txt)
                    print(f"[OK] {fn} -> {out} ({len(txt)} chars)")
                else:
                    bad_files.append(inp)

    # 结尾打印一个汇总，方便你一次性删除/替换
    if bad_files:
        print("\n[SKIPPED / BAD FILES]")
        for p in bad_files:
            print(" -", p)

def main():
    ap = argparse.ArgumentParser(description="将Word文档(.doc/.docx)转换为纯文本文件(.txt)")
    ap.add_argument("--in_dir", required=True, help="输入目录路径")
    ap.add_argument("--out_dir", required=True, help="输出目录路径")
    ap.add_argument("--recursive", action="store_true", default=True, help="递归处理子目录")
    args = ap.parse_args()
    
    process_directory(args.in_dir, args.out_dir, args.recursive)

if __name__ == "__main__":
    main()